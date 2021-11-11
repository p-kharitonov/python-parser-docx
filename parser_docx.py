from docx import Document

class ParserDocx:
  def __init__(self, path):
    self.path = path
    self.document = Document(path)

  def get_text_list(self):
    result = []
    for paragraph in self.document.paragraphs:
      result.append(paragraph.text)
    result += self._get_text_from_table()
    return result

  def get_text(self):
    return "\n".join(self.get_text_list())

  def _get_text_from_table(self, tables=None) -> list:
    result = []
    if tables is None:
      tables = self.document.tables
    for table in tables:
      for row in table.rows:
        for cell in row.cells:
          if cell.tables:
            result += self._get_text_from_table(cell.tables)
          for paragraph in cell.paragraphs:
            result.append(paragraph.text)
    return result


